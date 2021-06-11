import docx
from docx.shared import RGBColor, Pt
import glob
import re
import sys


def get_ans(file):
    f = open(file, "r")
    ans = f.read()
    regex = re.compile('[^A-Z]')
    ans = regex.sub('', ans)
    print(ans)
    # ans = [ord(char) - 64 for char in list(ans)]
    ans = list(ans)
    print(ans)
    return ans


def create_word(ans = False):
    doc = docx.Document()
    for index in range(len(images)):
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(images[index])

        if(ans):
            r.add_text(ans[index])
            r.font.color.rgb = RGBColor(255, 0, 0)
            r.font.size = Pt(18)
    return doc

if __name__ == '__main__':
    images = glob.glob("img/*.png")
    print(images)

    if len(sys.argv) == 2:
        ans = get_ans(sys.argv[1])
        doc = create_word(ans)
    else:
        doc = create_word()

    doc.save('Output.docx')
    print("\nFinish !!!")